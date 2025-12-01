from docx import Document;
from docx.shared import Pt, RGBColor, Inches;

def writeResponseWithReasoning(question, generalInfo, databaseInfo):
    return runAgent(
        "RFP Response Writer",
        """
You are working at Prompta AI, a company that uses AI technology to analyse and improve business practises of partnering organisations who are mostly Canadian government agencies and corporations.
        """,
        """
You are an experienced RFP response writer at Prompta AI, specialising in reading complex RFPs from corporations and government agencies and replying to them. Your current job is to write the response to a RFP proposal question received by Prompta AI using information from the company database.
        """,
        """
You will be provided with 3 pieces of information
A. the RFP question that you are to respond to, along with further requirements about how it should be answered
B. a summary of what the RFP is about to help you decide which points in the company database to focus on in your response
C. a bullet point list of information about Prompta AI. ALL information about Prompta AI included in your response MUST come from this database. IT IS STRICTLY FORBIDDEN to make up information.

You should follow this reasoning process
1. interpret what the question means and rephrase the requirements to make them applicable (for example rephrase a requirement about page count into one about word count)
2. interpret what the project that the client is sending RFPs for is about, and use this information to infer what technologies, processes, and services clients typically expect Prompta AI to provide
3. read through the company database information and pick out points that are relevant, also make sure to indicate which points you should emphasise
4. in a new line, write STARTING_RESPONSE to mark the start of your final answer
5. write your response to the RFP question using information that you have identified as relevant in step 3, make sure that you emphasise the important parts. Your response should be written clearly and in a tone appropriate for direct inclusion in Prompta AI's RFP response. After completing this step, DO NOT say anything else.

The information is provided below.

information A:
{question}

information B:
{generalInfo}

information C:
{databaseInfo}
        """,
        """
1 short paragraph each for reasoning steps 1 to 3, followed by the text STARTING_RESPONSE on a new line, followed by the actual response to the RFP question. The response must be presentable and be written in an appropriate tone for RFP response proposals. Your response must only use information identified in reasoning step 3, DO NOT MAKE UP INFORMATION.
        """,
        {
            "question": question,
            "generalInfo": generalInfo,
            "databaseInfo": databaseInfo
        },
        False
    );

def parseResponseWritten(response):
    index = response.find("STARTING_RESPONSE");
    if(index == -1 or index + 18 >= len(response)):
        return "";
    return response[index + 18::];

def DLL_EXPORT_writeResponse(question, generalInfo, databaseInfo):
    attemptCount = 0;
    while(attemptCount < 5):
        parsedResponse = parseResponseWritten(writeResponseWithReasoning(question, generalInfo, databaseInfo));
        if(parsedResponse != ""):
            return parsedResponse;
    return "";

def DLL_EXPORT_enhanceResponse(response, instructions, questionNumber, question, infoDatabase):
    improvedResponse = "";
    if(questionNumber == -1):
        improvedResponse = runAgent(
            "RFP Response Writer",
            """
You are working at Prompta AI, a company that uses AI technology to analyse and improve business practises of partnering organisations who are mostly Canadian government agencies and corporations.
            """,
            """
You are an experienced RFP response writer at Prompta AI, specialising in reading complex RFPs from corporations and government agencies and replying to them. Your current job is to improve the draft response of a RFP question.
            """,
            """
You will be provided with the existing draft response and some instructions. You are to improve the draft response according to the instructions, but you must not change the information presented. You are allowed and encouraged to rephrase or combine points if necessary, but you must not add information or remove information.

Ensure that the improved response is grammatically correct in Canadian English and is written in a tone suitable for formal communications between Canadian corporations.

The draft response and instructions are provided below.

Draft response:
{draft}

Instructions:
{instructions}
            """,
            """
Output only the improved version of the response. You must not include anything else such as "here is your response" or "how may I help you". Your output will be directly used by the Prompta AI. Remember that your improved version must convey the exact same information as the draft provided.
            """,
            {
                "draft": response,
                "instructions": instructions
            },
            False
        );
    else:
        improvedResponse = runAgent(
            "RFP Response Writer",
            """
You are working at Prompta AI, a company that uses AI technology to analyse and improve business practises of partnering organisations who are mostly Canadian government agencies and corporations.
            """,
            """
You are an experienced RFP response writer at Prompta AI, specialising in reading complex RFPs from corporations and government agencies and replying to them. Your current job is to improve the draft response of a RFP question.
            """,
            """
You will be provided with the existing draft response, some instructions, the question, and a database containing information about Prompta AI. You are to improve the draft response according to the instructions.

You should follow this reasoning process
1. read through the draft response and the question. You may assume that the information in the draft response is reliable. Analyse if the draft response is convincing and if it answers the question well.
2. based on your answer to step 1, suggest how can the draft response be improved (but do not rewrite it yet) so that it can become more convincing and relevant. You must make sure that you follow the instructions provided below when making these suggestions. Follow the instructions instead of questions requirements in case of conflict.
3. based on your answer to step 1 and 2, and with use of the information database, pick out points that you really have to include in a improved version of the response. Note that if the instructions do not explicitly mention adding information, this should be kept minimal or none at all.
4. in a new line, write STARTING_RESPONSE to mark the start of your final answer
5. write your improved response using information that you have identified as relevant in step 3, make sure that you emphasise the important parts. Ensure that the improved response is grammatically correct in Canadian English and is written in a tone suitable for formal communications between Canadian corporations. Ensure that instructions below and your suggestions from step 2 are followed. After completing this step, DO NOT say anything else.

Everything is provided below.

Draft response:
{draft}

Instructions:
{instructions}

Question with requirements:
{question}

Information Database:
{infoDatabase}
            """,
            """
1 short paragraph each for reasoning steps 1 to 3, followed by the text STARTING_RESPONSE on a new line, followed by the improved response. The improved response must be presentable and be written in an appropriate tone for RFP response proposals. DO NOT MAKE UP INFORMATION. You must not include anything else such as "here is your response" or "how may I help you". Your output will be directly used by Prompta AI.
            """,
            {
                "draft": response,
                "instructions": instructions,
                "question": question,
                "infoDatabase": infoDatabase
            },
            False
        );
        improvedResponse = parseResponseWritten(improvedResponse);
    return improvedResponse;

def createDocumentStructure(info):
    global API_MSG_generateDocxMessage;
    API_MSG_generateDocxMessage = "formatting document...";
    return runAgent(
        "RFP Response Writer",
        """
You are working at Prompta AI, a company that uses AI technology to analyse and improve business practises of partnering organisations who are mostly Canadian government agencies and corporations.
        """,
        """
You are an experienced RFP response writer at Prompta AI, specialising in reading complex RFPs from corporations and government agencies and replying to them. Your current job is to assemble responses to RFP questions written by different teams into a single document.
        """,
        """
You will be provided with all the questions from the RFPs each marked with QUESTION and the responses written marked with RESPONSE immediately after each question. The responses were written by different teams, so the formatting style is very likely different. You should reformat them to keep a consistent style.

Your output should be a single document containing all the responses in the same order as they were given to you. This document should be directly presentable in a Canadian professional / business setting. The response to each question should start with the identification code or index number of the question in the header, if any were provided.

You should structure your document properly. Each response should start with a header, marked with @HEADER at the start of the line containing the header. To indicate important sections within a response, mark it with @SUBHEADER at the start of its line. Your output will be directly used to generate the document, so do not include anything other than the reformatted document.

questions and responses:

{info}
        """,
        """
Only the reformatted document with headers and sub headers marked appropriately. You must not include anything else such as "here is your response" or "how may I help you". Your output will be directly used by Prompta AI.
        """,
        {"info": info},
        False
    );

def deduplicateFinalDocument(info, useEnhance):
    global API_MSG_generateDocxMessage;
    API_MSG_generateDocxMessage = "removing duplicate information...";
    if(useEnhance == 0):
        return info;
    return runAgent(
        "RFP Response Writer",
        """
You are working at Prompta AI, a company that uses AI technology to analyse and improve business practises of partnering organisations who are mostly Canadian government agencies and corporations.
        """,
        """
You are an experienced RFP response writer at Prompta AI, specialising in reading complex RFPs from corporations and government agencies and replying to them. Your current job is to remove duplicate information from a RFP response proposal.
        """,
        """
You are provided with information from a RFP proposal document below. Some responses may contain duplicate information. You are to keep only that information in the most relevant response and remove it from anywhere else. The @HEADER and @SUBHEADER are important for formatting and should not be touched. Your output will be directly used to generate the document, so do not include anything other than the reformatted document.

proposal document:

{info}
        """,
        """
Only the edited document. You must not include anything else such as "here is your response" or "how may I help you". Your output will be directly used by Prompta AI.
        """,
        {"info": info},
        False
    );

def generateDOCX(contents):
    global API_MSG_generateDocxMessage;
    API_MSG_generateDocxMessage = "generating DOCX file...";
    document = Document();
    infoList = [];
    while(True):
        if(contents == ""):
            break;
        if(contents.find("\n") == -1):
            infoList.append(contents);
            contents = "";
            break;
        index = contents.find("\n");
        infoList.append(contents[0:index:]);
        contents = contents[index+1::];
    for i in infoList:
        infoType = 0; # 0 - normal, 1 - header, 2 - subheader
        if(i[0:7:] == "@HEADER"):
            i = i[8::];
            infoType = 1;
        elif(i[0:10:] == "@SUBHEADER"):
            i = i[11::];
            infoType = 2;
        textElement = document.add_paragraph().add_run(i);
        textElement.font.name = "Arial";
        textElement.font.size = Pt(11);
        if(infoType != 0):
            textElement.font.bold = True;
        if(infoType == 1):
            textElement.font.color.rgb = RGBColor(31, 72, 124);
    stringStream = io.BytesIO();
    document.save(stringStream);
    stringStream.seek(0);
    return bytearray(stringStream.getvalue());
