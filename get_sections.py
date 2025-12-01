from collections import namedtuple;
from docx import Document;
import fitz;
from docx.oxml.ns import qn;
from lxml import etree;
import numpy as np;
from PIL import Image;
import io;
import easyocr;
import zipfile;

Span = namedtuple("Span", ["text", "style", "xCoord", "maybeHeader"]);
Section = namedtuple("Section", ["header", "text"]);

def cleanSpace(text):
    outputText = "";
    for c in text:
        if(c != '\n' and c != '\r' and c != ' '):
            outputText += c;
    return outputText;

def cleanLeadingTrailingSpace(text):
    firstNonSpace = 0;
    for i in range(len(text)):
        if(text[i] != '\n' and text[i] != '\r' and text[i] != ' '):
            firstNonSpace = i;
            break;
    text = text[firstNonSpace::];
    for i in range(len(text)):
        if(text[len(text) - 1 - i] != '\n' and text[len(text) - 1 - i] != '\r' and text[len(text) - 1 - i] != ' '):
            firstNonSpace = len(text) - 1 - i;
            break;
    text = text[:firstNonSpace + 1:];
    return text;

def getTextFromImage(blob):
    if(blob == None):
        return "";
    img = np.array(Image.open(io.BytesIO(blob)).convert("RGB"));
    result = easyocr.Reader(["en"], gpu=True).readtext(img, detail = 0);
    output = "\n";
    for i in range(len(result)):
        output += (result[i] + '\n');
    return ("\n[below is text found in a image, ignore if nonsensical]" + output + "[end image text block]\n");

def getImageInDOCX(run):
    blob = None;
    if(run.element == None):
        return None;
    draw = run.element.xpath(".//w:drawing");
    if(draw == None):
        return None;
    for p in draw:
        xmlNamespace = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture"
        };
        realXMLElement = etree.fromstring(p.xml.encode("utf-8"));
        bl = realXMLElement.xpath(".//a:blip", namespaces = xmlNamespace);
        if(bl != None):
            for b in bl:
                embedID = b.get(qn("r:embed"));
                if(embedID != None and embedID in run.part.related_parts):
                    blob = run.part.related_parts[embedID].blob;
    return blob;

def getTextFromDOCX(fileBlob):
    global API_MSG_projectCreationMessage;
    global API_MSG_uploadResponseMessage;
    output = [];
    document = Document(io.BytesIO(fileBlob));
    paragraphCount = 0;
    tableCount = 0;
    progressCounter = 0;
    for child in document.element.body:
        API_MSG_projectCreationMessage = "reading DOC / DOCX file... " + str(int(100.0 * float(progressCounter) / float(len(document.element.body)))) + "%";
        API_MSG_uploadResponseMessage = "reading DOC / DOCX file... " + str(int(100.0 * float(progressCounter) / float(len(document.element.body)))) + "%";
        progressCounter += 1;
        if(child.tag == qn("w:p")):
            paragraph = document.paragraphs[paragraphCount];
            paragraphCount += 1;
            for run in paragraph.runs:
                if(cleanLeadingTrailingSpace(run.text) == ""):
                    imgText = getTextFromImage(getImageInDOCX((run)));
                    output.append(Span(imgText, 0, 0.0, False));
                    continue;
                styleFlag = 0;
                if(run.bold == True):
                    styleFlag = 16; # 16 is bold flag, this is if word is bold
                elif(run.style != None and run.style.font.bold == True):
                    styleFlag = 16; # inherit from run styling
                elif(paragraph.style != None and paragraph.style.font.bold == True):
                    styleFlag = 16; # inherit from paragraph styling
                elif(document.styles["Normal"].font.bold == True):
                    styleFlag = 16; # check DOCX default style presets
                output.append(Span(cleanLeadingTrailingSpace(run.text), styleFlag, 0.0, True));
        elif(child.tag == qn("w:tbl")):
            table = document.tables[tableCount];
            tableCount += 1;
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if(cleanLeadingTrailingSpace(run.text) == ""):
                                imgText = getTextFromImage(getImageInDOCX((run)));
                                output.append(Span(imgText, 0, 0.0, False));
                                continue;
                            styleFlag = 0;
                            if(run.bold == True):
                                styleFlag = 16; # 16 is bold flag, this is if word is bold
                            elif(run.style != None and run.style.font.bold == True):
                                styleFlag = 16; # inherit from run styling
                            elif(paragraph.style != None and paragraph.style.font.bold == True):
                                styleFlag = 16; # inherit from paragraph styling
                            elif(document.styles["Normal"].font.bold == True):
                                styleFlag = 16; # check DOCX default style presets
                            output.append(Span(cleanLeadingTrailingSpace(run.text), styleFlag, 0.0, True));
    return output;

def getTextFromPDF(fileBlob):
    global API_MSG_projectCreationMessage;
    global API_MSG_uploadResponseMessage;
    output = [];
    document = fitz.open(stream = bytes(fileBlob), filetype = "pdf");
    progressCounter = 0;
    for page in document:
        API_MSG_projectCreationMessage = "reading PDF file... " + str(int(100.0 * float(progressCounter) / float(len(document)))) + "%";
        API_MSG_uploadResponseMessage = "reading PDF file... " + str(int(100.0 * float(progressCounter) / float(len(document)))) + "%";
        progressCounter += 1;
        textInfo = page.get_text("dict")["blocks"];
        for block in textInfo:
            if("lines" in block):
                for line in block["lines"]:
                    for span in line["spans"]:
                        if(span["bbox"][1] > 0.9 * page.rect.height):
                            continue;
                        if(cleanSpace(span["text"]) == ""):
                            continue;
                        if(len(line["spans"]) > 1):
                            output.append(Span(cleanLeadingTrailingSpace(span["text"]), span["flags"], span["bbox"][0], False));
                        else:
                            output.append(Span(cleanLeadingTrailingSpace(span["text"]), span["flags"], span["bbox"][0], True));
            elif(block["type"] == 1):
                xref = block.get("xref");
                blob = None;
                if(xref != None):
                    blob = document.extract_image(xref)["image"];
                else:
                    blob = block.get("image");
                imgText = getTextFromImage(blob);
                output.append(Span(imgText, 0, 0.0, False));
    return output;

def textIsHeader(span):
    if((span.style & 16) != 0 and span.maybeHeader == True):
        return True;
    return False;

def textIsBold(span):
    if((span.style & 16) != 0):
        return True;
    return False;

def isTopLevelSection(header):
    for i in range(len(header)):
        if((header[i] >= 'a' and header[i] <= 'z') or (header[i] >= 'A' and header[i] <= 'Z') or (header[i] >= '0' and header[i] <= '9')):
            i -= 1;
            break;
    header = header[i+1::];
    if(header == ""):
        return False;
    for i in range(len(header)):
        if((header[i] < 'a' or header[i] > 'z') and (header[i] < 'A' or header[i] > 'Z') and (header[i] < '0' or header[i] > '9')):
            break;
    header = header[:i:];
    if(len(header) > 6 and header[:6:].upper() == "APPEND"):
        return True;
    try:
        j = int(header);
        return True;
    except:
        return False;

CACHE_documentSections = [];
def getTextFromZIP(fileBlob, isReceiving):
    global CACHE_documentSections;
    zipFileManager = zipfile.ZipFile(io.BytesIO(fileBlob), "r");
    output = [];
    for i in zipFileManager.namelist():
        file = zipFileManager.open(i);
        extractedFileBlob = file.read();
        file.close();
        if(len(i) > 4 and i[len(i)-4::].lower() == ".pdf"):
            getSectionsFromDocument(extractedFileBlob, 1, isReceiving);
            output.extend(CACHE_documentSections.copy());
        elif(len(i) > 5 and i[len(i)-5::].lower() == ".docx"):
            getSectionsFromDocument(extractedFileBlob, 2, isReceiving);
            output.extend(CACHE_documentSections.copy());
        elif(len(i) > 4 and i[len(i)-5::].lower() == ".doc"):
            getSectionsFromDocument(extractedFileBlob, 3, isReceiving);
            output.extend(CACHE_documentSections.copy());
    zipFileManager.close();
    return output;

def getSectionsFromDocument(fileBlob, fileType, isReceiving):
    global API_MSG_projectCreationMessage;
    global API_MSG_uploadResponseMessage;
    global CACHE_documentSections;
    sections = [];
    currentSection = Section("", "");
    sourceText = [];
    if(fileType == 1):
        sourceText = getTextFromPDF(fileBlob);
    elif(fileType == 2 or fileType == 3):
        sourceText = getTextFromDOCX(fileBlob);
    elif(fileType == 4):
        CACHE_documentSections = getTextFromZIP(fileBlob, isReceiving);
        return;
    API_MSG_projectCreationMessage = "reconstructing document formatting...";
    API_MSG_uploadResponseMessage = "reconstructing document formatting...";
    leftAlign = float("inf");
    noLongerHeader = [];
    for i in range(len(sourceText)):
        if(i == 0):
            continue;
        if(textIsHeader(sourceText[i]) == True and textIsHeader(sourceText[i-1]) == True):
            noLongerHeader.append(i);
    for i in noLongerHeader:
        sourceText[i] = Span(sourceText[i].text, sourceText[i].style, sourceText[i].xCoord, False);
    for span in sourceText:
        if(textIsHeader(span) == True):
            if(span.xCoord < leftAlign):
                leftAlign = span.xCoord;
    inHeader = True;
    for span in sourceText:
        if(textIsHeader(span) == True and span.xCoord <= 1.1 * leftAlign):
            if(currentSection.header != ""):
                sections.append(currentSection);
            currentSection = Section(span.text, "");
            inHeader = True;
        else:
            if(inHeader == True):
                if(textIsBold(span) == True):
                    currentSection = Section(currentSection.header + " " + span.text, "");
                else:
                    currentSection = Section(currentSection.header, span.text);
                    inHeader = False;
            else:
                currentSection = Section(currentSection.header, currentSection.text + " " + span.text);
    if(currentSection.header != ""):
        sections.append(currentSection);
    mergedSections = [];
    thisSection = Section("", "");
    for section in sections:
        if(isReceiving == False or isTopLevelSection(section.header) == True):
            if(thisSection.header != ""):
                mergedSections.append(thisSection);
                thisSection = Section("", "");
            thisSection = section;
        else:
            if(thisSection.header == ""):
                thisSection = section;
            else:
                thisSection = Section(thisSection.header, thisSection.text + "\n\n" + section.header + "\n" + section.text)
    if(thisSection.header != ""):
        mergedSections.append(thisSection);
    CACHE_documentSections = mergedSections;
